"""Multi-format document parser.

Supports: .pdf, .docx, .ipynb, .md, .txt
Extracts raw text, sections, and raw equation strings.
"""
import re
from pathlib import Path
from typing import Dict, List, Tuple

from app.models.validation import ParsedDocument


def parse_document(file_path: Path) -> ParsedDocument:
    suffix = file_path.suffix.lower()
    dispatch = {
        ".pdf": _parse_pdf,
        ".docx": _parse_docx,
        ".ipynb": _parse_notebook,
        ".md": _parse_text,
        ".txt": _parse_text,
    }
    parser = dispatch.get(suffix)
    if not parser:
        raise ValueError(f"Formato não suportado: {suffix}")

    raw_text, sections, equations = parser(file_path)
    return ParsedDocument(
        filename=file_path.name,
        format=suffix.lstrip("."),
        raw_text=raw_text,
        sections=sections,
        equations_raw=equations,
        metadata={"path": str(file_path), "size_bytes": file_path.stat().st_size},
    )


def _extract_latex(text: str) -> List[str]:
    """Extract LaTeX math blocks: $...$, $$...$$, \[...\], \begin{equation}..."""
    patterns = [
        r"\$\$(.+?)\$\$",
        r"\$([^$\n]+?)\$",
        r"\\\[(.+?)\\\]",
        r"\\begin\{equation\*?\}(.+?)\\end\{equation\*?\}",
        r"\\begin\{align\*?\}(.+?)\\end\{align\*?\}",
    ]
    equations = []
    for p in patterns:
        for m in re.finditer(p, text, re.DOTALL):
            eq = m.group(1).strip()
            if eq:
                equations.append(eq)
    return equations


def _split_sections(text: str) -> Dict[str, str]:
    """Split markdown-style headings into sections dict."""
    sections: Dict[str, str] = {}
    current = "preamble"
    buf: List[str] = []
    for line in text.splitlines():
        m = re.match(r"^(#{1,4})\s+(.+)", line)
        if m:
            if buf:
                sections[current] = "\n".join(buf).strip()
            current = m.group(2).strip()
            buf = []
        else:
            buf.append(line)
    if buf:
        sections[current] = "\n".join(buf).strip()
    return sections


def _parse_text(path: Path) -> Tuple[str, Dict, List]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return text, _split_sections(text), _extract_latex(text)


def _parse_pdf(path: Path) -> Tuple[str, Dict, List]:
    import fitz  # PyMuPDF
    doc = fitz.open(str(path))
    pages = [page.get_text() for page in doc]
    text = "\n".join(pages)
    doc.close()
    return text, _split_sections(text), _extract_latex(text)


def _parse_docx(path: Path) -> Tuple[str, Dict, List]:
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(path))
    lines = []
    equations = []

    for para in doc.paragraphs:
        # Check for OMML equations embedded in paragraph
        omml_elems = para._element.findall(
            f".//{{{qn('m:oMath')}}}", para._element.nsmap
        ) if hasattr(para._element, "nsmap") else []
        if omml_elems:
            # Store raw XML representation as placeholder
            equations.append(f"[OMML_EQUATION: {para.text or 'inline eq'}]")

        heading_style = para.style.name.lower() if para.style else ""
        if "heading" in heading_style:
            level = "".join(filter(str.isdigit, heading_style)) or "1"
            lines.append(f"{'#' * int(level)} {para.text}")
        else:
            lines.append(para.text)

    text = "\n".join(lines)
    equations += _extract_latex(text)
    return text, _split_sections(text), equations


def _parse_notebook(path: Path) -> Tuple[str, Dict, List]:
    import nbformat
    nb = nbformat.read(str(path), as_version=4)
    parts = []
    equations = []

    for cell in nb.cells:
        if cell.cell_type == "markdown":
            parts.append(cell.source)
            equations += _extract_latex(cell.source)
        elif cell.cell_type == "code":
            parts.append(f"```python\n{cell.source}\n```")
            # capture outputs
            for out in cell.get("outputs", []):
                if out.get("output_type") == "stream":
                    parts.append(out.get("text", ""))
                elif "text/plain" in out.get("data", {}):
                    parts.append(out["data"]["text/plain"])

    text = "\n\n".join(parts)
    return text, _split_sections(text), equations
