"""Tests for all document, excel, and code parsers."""
import json
import pytest
from pathlib import Path
from unittest.mock import MagicMock, patch

from app.services.parsers.document import (
    parse_document, _extract_latex, _split_sections,
    _parse_text, _parse_pdf, _parse_docx, _parse_notebook,
)
from app.services.parsers.excel import parse_excel, excel_to_text
from app.services.parsers.code import parse_code_files, _detect_language
from app.models.validation import ParsedDocument, ParsedCode


# ── _extract_latex ────────────────────────────────────────────────────────────

class TestExtractLatex:
    def test_double_dollar(self):
        eqs = _extract_latex("$$PV = N \\cdot e^{-r}$$")
        assert any("PV" in e for e in eqs)

    def test_single_dollar(self):
        eqs = _extract_latex("The formula $r = 0.12$ applies.")
        assert any("r = 0.12" in e for e in eqs)

    def test_bracket_syntax(self):
        eqs = _extract_latex(r"\[F = S \cdot e^{rT}\]")
        assert any("F = S" in e for e in eqs)

    def test_begin_equation(self):
        eqs = _extract_latex(r"\begin{equation}C = S N(d_1)\end{equation}")
        assert any("C = S" in e for e in eqs)

    def test_begin_align(self):
        eqs = _extract_latex(r"\begin{align}d_1 &= \frac{\ln(S/K)}{v\sqrt{T}}\end{align}")
        assert any("d_1" in e for e in eqs)

    def test_no_equations(self):
        eqs = _extract_latex("Plain text without any math.")
        assert eqs == []

    def test_multiple_equations(self):
        text = "First $a+b$ then $$c-d$$"
        eqs = _extract_latex(text)
        assert len(eqs) >= 2

    def test_empty_equation_skipped(self):
        eqs = _extract_latex("$$$$")  # empty double dollar
        assert all(e.strip() for e in eqs)


# ── _split_sections ───────────────────────────────────────────────────────────

class TestSplitSections:
    def test_single_heading(self):
        text = "# Title\nsome content here"
        sections = _split_sections(text)
        assert "Title" in sections
        assert "content" in sections["Title"]

    def test_multiple_headings(self):
        text = "# Sec1\ntext1\n## Sec2\ntext2\n### Sec3\ntext3"
        sections = _split_sections(text)
        assert "Sec1" in sections
        assert "Sec2" in sections
        assert "Sec3" in sections

    def test_preamble_before_first_heading(self):
        text = "intro text\n# Section\ncontent"
        sections = _split_sections(text)
        assert "preamble" in sections
        assert "intro" in sections["preamble"]

    def test_no_headings(self):
        text = "just plain text without headings"
        sections = _split_sections(text)
        assert "preamble" in sections

    def test_empty_text(self):
        sections = _split_sections("")
        assert isinstance(sections, dict)


# ── _parse_text ───────────────────────────────────────────────────────────────

class TestParseText:
    def test_markdown_file(self, tmp_path):
        md = tmp_path / "model.md"
        md.write_text("# Modelo\n\nEquação: $$PV = N$$\n", encoding="utf-8")
        text, sections, eqs = _parse_text(md)
        assert "Modelo" in text
        assert len(eqs) > 0

    def test_txt_file(self, tmp_path):
        txt = tmp_path / "model.txt"
        txt.write_text("Simple model description.", encoding="utf-8")
        text, sections, eqs = _parse_text(txt)
        assert "Simple model" in text


# ── _parse_pdf ────────────────────────────────────────────────────────────────

class TestParsePDF:
    def test_extracts_text_from_pages(self, tmp_path):
        pdf_path = tmp_path / "model.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 fake")

        mock_page = MagicMock()
        mock_page.get_text.return_value = "# Model\n\n$$PV = N \\cdot e^{-r}$$"
        mock_doc = MagicMock()
        mock_doc.__iter__ = MagicMock(return_value=iter([mock_page]))
        mock_doc.close = MagicMock()

        with patch("fitz.open", return_value=mock_doc):
            text, sections, eqs = _parse_pdf(pdf_path)

        assert "Model" in text
        assert len(eqs) > 0
        mock_doc.close.assert_called_once()

    def test_multiple_pages(self, tmp_path):
        pdf_path = tmp_path / "m.pdf"
        pdf_path.write_bytes(b"fake")

        pages = [MagicMock(), MagicMock()]
        pages[0].get_text.return_value = "# Section 1\npage one"
        pages[1].get_text.return_value = "## Section 2\npage two"
        mock_doc = MagicMock()
        mock_doc.__iter__ = MagicMock(return_value=iter(pages))
        mock_doc.close = MagicMock()

        with patch("fitz.open", return_value=mock_doc):
            text, _, _ = _parse_pdf(pdf_path)

        assert "page one" in text
        assert "page two" in text


# ── _parse_docx ───────────────────────────────────────────────────────────────

class TestParseDocx:
    def test_real_docx_with_heading(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document as DocxDoc
        doc = DocxDoc()
        doc.add_heading("Model Title", level=1)
        doc.add_paragraph("This is the model description.")
        path = tmp_path / "model.docx"
        doc.save(str(path))

        text, sections, eqs = _parse_docx(path)
        assert "Model Title" in text
        assert "model description" in text

    def test_docx_with_latex_in_paragraph(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document as DocxDoc
        doc = DocxDoc()
        doc.add_paragraph(r"Formula: $PV = N \cdot e^{-r}$")
        path = tmp_path / "model2.docx"
        doc.save(str(path))

        text, sections, eqs = _parse_docx(path)
        assert len(eqs) > 0

    def test_docx_para_without_style(self, tmp_path):
        """Covers the `if para.style else ""` branch."""
        mock_para = MagicMock()
        mock_para.style = None
        mock_para.text = "no style paragraph"
        mock_para._element = MagicMock(spec=[])  # no nsmap attr

        mock_doc = MagicMock()
        mock_doc.paragraphs = [mock_para]

        with patch("docx.Document", return_value=mock_doc):
            with patch("docx.oxml.ns.qn", return_value="m:oMath"):
                text, _, _ = _parse_docx(tmp_path / "x.docx")

        assert "no style paragraph" in text

    def test_docx_para_with_nsmap_and_omml(self, tmp_path):
        """Covers the OMML equation detection branch."""
        mock_elem = MagicMock()
        mock_elem.nsmap = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}
        mock_elem.findall = MagicMock(return_value=["eq_element"])

        mock_para = MagicMock()
        mock_para.style = None
        mock_para.text = ""
        mock_para._element = mock_elem

        mock_doc = MagicMock()
        mock_doc.paragraphs = [mock_para]

        with patch("docx.Document", return_value=mock_doc):
            with patch("docx.oxml.ns.qn", return_value="m:oMath"):
                text, _, eqs = _parse_docx(tmp_path / "x.docx")

        assert any("OMML" in e for e in eqs)


# ── _parse_notebook ───────────────────────────────────────────────────────────

class TestParseNotebook:
    def test_markdown_and_code_cells(self, tmp_path):
        nbformat = pytest.importorskip("nbformat")
        nb = nbformat.v4.new_notebook()
        nb.cells = [
            nbformat.v4.new_markdown_cell("# Intro\n$$PV = N$$"),
            nbformat.v4.new_code_cell("print('hello')"),
        ]
        path = tmp_path / "model.ipynb"
        nbformat.write(nb, str(path))

        text, sections, eqs = _parse_notebook(path)
        assert "Intro" in text
        assert "print" in text
        assert len(eqs) > 0

    def test_stream_output(self, tmp_path):
        nbformat = pytest.importorskip("nbformat")
        nb = nbformat.v4.new_notebook()
        code_cell = nbformat.v4.new_code_cell("print('result')")
        code_cell.outputs = [{"output_type": "stream", "text": "result\n", "name": "stdout"}]
        nb.cells = [code_cell]
        path = tmp_path / "nb.ipynb"
        nbformat.write(nb, str(path))

        text, _, _ = _parse_notebook(path)
        assert "result" in text

    def test_text_plain_output(self, tmp_path):
        nbformat = pytest.importorskip("nbformat")
        nb = nbformat.v4.new_notebook()
        code_cell = nbformat.v4.new_code_cell("1+1")
        code_cell.outputs = [{"output_type": "execute_result", "data": {"text/plain": "2"}, "metadata": {}, "execution_count": 1}]
        nb.cells = [code_cell]
        path = tmp_path / "nb2.ipynb"
        nbformat.write(nb, str(path))

        text, _, _ = _parse_notebook(path)
        assert "2" in text


# ── parse_document dispatch ───────────────────────────────────────────────────

class TestParseDocument:
    def test_dispatches_markdown(self, tmp_path):
        md = tmp_path / "model.md"
        md.write_text("# Model\nContent", encoding="utf-8")
        doc = parse_document(md)
        assert isinstance(doc, ParsedDocument)
        assert doc.format == "md"
        assert doc.filename == "model.md"

    def test_dispatches_txt(self, tmp_path):
        txt = tmp_path / "model.txt"
        txt.write_text("Model content", encoding="utf-8")
        doc = parse_document(txt)
        assert doc.format == "txt"

    def test_dispatches_pdf(self, tmp_path):
        pdf = tmp_path / "model.pdf"
        pdf.write_bytes(b"fake")
        mock_page = MagicMock()
        mock_page.get_text.return_value = "content"
        mock_doc = MagicMock()
        mock_doc.__iter__ = MagicMock(return_value=iter([mock_page]))
        mock_doc.close = MagicMock()

        with patch("fitz.open", return_value=mock_doc):
            doc = parse_document(pdf)

        assert doc.format == "pdf"

    def test_dispatches_notebook(self, tmp_path):
        nbformat = pytest.importorskip("nbformat")
        nb = nbformat.v4.new_notebook()
        nb.cells = [nbformat.v4.new_markdown_cell("# Test")]
        path = tmp_path / "nb.ipynb"
        nbformat.write(nb, str(path))
        doc = parse_document(path)
        assert doc.format == "ipynb"

    def test_unsupported_extension_raises(self, tmp_path):
        bad = tmp_path / "model.xyz"
        bad.write_text("data")
        with pytest.raises(ValueError, match="Formato não suportado"):
            parse_document(bad)

    def test_metadata_contains_size(self, tmp_path):
        md = tmp_path / "m.md"
        md.write_text("# Test\nContent here", encoding="utf-8")
        doc = parse_document(md)
        assert "size_bytes" in doc.metadata
        assert doc.metadata["size_bytes"] > 0


# ── Excel parser ──────────────────────────────────────────────────────────────

class TestExcelParser:
    def _create_excel(self, tmp_path, with_formula=True, with_named_range=False):
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Pricing"
        ws["A1"] = "Factor"
        ws["B1"] = "Value"
        ws["A2"] = "taxa_di"
        ws["B2"] = 0.12
        if with_formula:
            ws["C2"] = "=B2*1000000"
        path = tmp_path / "model.xlsx"
        wb.save(str(path))
        return path

    def test_parse_excel_returns_sheets(self, tmp_path):
        path = self._create_excel(tmp_path)
        sheets, named = parse_excel(path)
        assert len(sheets) == 1
        assert sheets[0]["sheet_name"] == "Pricing"

    def test_parse_excel_detects_formulas(self, tmp_path):
        path = self._create_excel(tmp_path, with_formula=True)
        sheets, named = parse_excel(path)
        assert sheets[0]["formula_count"] > 0
        assert any("=B2" in f for f in sheets[0]["formulas_sample"])

    def test_parse_excel_dimensions(self, tmp_path):
        path = self._create_excel(tmp_path)
        sheets, _ = parse_excel(path)
        assert "x" in sheets[0]["dimensions"]

    def test_parse_excel_named_ranges_empty(self, tmp_path):
        path = self._create_excel(tmp_path)
        _, named = parse_excel(path)
        assert isinstance(named, dict)

    def test_excel_to_text_returns_string(self, tmp_path):
        path = self._create_excel(tmp_path)
        text = excel_to_text(path)
        assert isinstance(text, str)
        assert "Pricing" in text
        assert "taxa_di" in text or "Fórmulas" in text

    def test_excel_to_text_with_formulas(self, tmp_path):
        path = self._create_excel(tmp_path, with_formula=True)
        text = excel_to_text(path)
        assert "=B2" in text

    def test_excel_named_range_exception_handling(self, tmp_path):
        """Covers the except: pass branch in named_ranges parsing."""
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        path = tmp_path / "nb.xlsx"
        wb.save(str(path))

        bad_range = MagicMock()
        bad_range.attr_text = property(lambda self: (_ for _ in ()).throw(AttributeError("no attr")))

        with patch("openpyxl.load_workbook") as mock_load:
            mock_wb = MagicMock()
            mock_wb.sheetnames = ["Sheet1"]
            mock_ws = MagicMock()
            mock_ws.iter_rows.return_value = []
            mock_ws.max_row = 0
            mock_ws.max_column = 0
            mock_wb.__getitem__ = MagicMock(return_value=mock_ws)
            mock_wb.defined_names.items.return_value = [("bad_range", bad_range)]
            mock_wb.close = MagicMock()
            mock_load.return_value = mock_wb

            sheets, named = parse_excel(path)

        assert isinstance(named, dict)


# ── Code parser ───────────────────────────────────────────────────────────────

class TestCodeParser:
    def test_detect_language(self):
        assert _detect_language(".py") == "python"
        assert _detect_language(".m") == "matlab"
        assert _detect_language(".sql") == "sql"
        assert _detect_language(".r") == "r"
        assert _detect_language(".jl") == "julia"
        assert _detect_language(".cpp") == "cpp"
        assert _detect_language(".c") == "c"
        assert _detect_language(".f90") == "fortran"
        assert _detect_language(".xyz") == "unknown"

    def test_parse_python_file(self, tmp_path):
        py = tmp_path / "pricer.py"
        py.write_text("import numpy as np\n\nclass ModelPricer:\n    pass\n")
        result = parse_code_files([py])
        assert isinstance(result, ParsedCode)
        assert len(result.files) == 1
        assert result.files[0]["language"] == "python"
        assert "python" in result.language

    def test_parse_multiple_languages(self, tmp_path):
        py = tmp_path / "pricer.py"
        py.write_text("class P: pass")
        sql = tmp_path / "query.sql"
        sql.write_text("SELECT * FROM model")
        result = parse_code_files([py, sql])
        langs = result.language
        assert "python" in langs
        assert "sql" in langs

    def test_parse_excel_in_code_files(self, tmp_path):
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws["A1"] = "test"
        path = tmp_path / "model.xlsx"
        wb.save(str(path))
        result = parse_code_files([path])
        assert any(f["language"] == "excel" for f in result.files)

    def test_content_capped_at_50k(self, tmp_path):
        py = tmp_path / "large.py"
        py.write_text("x = 1\n" * 10_000)
        result = parse_code_files([py])
        assert len(result.files[0]["content"]) <= 50_001

    def test_empty_file_list(self):
        result = parse_code_files([])
        assert result.files == []
        assert result.language == "unknown"

    def test_summary_format(self, tmp_path):
        py = tmp_path / "a.py"
        py.write_text("pass")
        result = parse_code_files([py])
        assert "arquivo" in result.summary
        assert "python" in result.summary
